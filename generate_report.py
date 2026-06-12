# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(12)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("人工智能前沿技术")
run.font.size = Pt(22)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("大作业")
run.font.size = Pt(18)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("人工智能学院")
run.font.size = Pt(14)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("2026年6月")
run.font.size = Pt(14)

doc.add_page_break()

doc.add_heading("一、项目概述", level=1)

doc.add_heading("1. 项目名称与选题方向", level=2)
doc.add_paragraph("项目名称：沉浸式多角色互动小说引擎")
doc.add_paragraph("选题方向：方向四——复杂状态与记忆方向")

doc.add_heading("2. 业务场景与痛点分析", level=2)
doc.add_paragraph(
    "传统文字冒险游戏存在剧情固定、角色互动有限、缺乏个性化体验等痛点。"
    "玩家在体验过程中无法真正感受到自由选择的魅力，NPC的行为模式单一，"
    "对话内容缺乏深度和连贯性。本项目旨在利用大语言模型（LLM）的自然语言理解和生成能力，"
    "结合LangChain框架的工具调用、记忆管理和RAG检索技术，打造一个能够动态生成剧情、"
    "自主决策、具备长短期记忆的AI驱动互动小说引擎。"
)

doc.add_heading("3. 核心功能清单", level=2)
doc.add_paragraph("功能点 1：支持多世界主题选择（暗影纪元/仙途），每个世界有独立的世界观设定和NPC角色体系")
doc.add_paragraph("功能点 2：基于ChatPromptTemplate注入复杂角色背景和性格特征，AI游戏大师自主推进剧情")
doc.add_paragraph("功能点 3：长短期记忆结合——ConversationBufferWindowMemory记录近期对话，ChromaDB向量数据库存储世界观和历史线索")
doc.add_paragraph("功能点 4：自主工具调用——探索、战斗、获取线索、查看状态、休息等工具，AI根据剧情需要自动触发")
doc.add_paragraph("功能点 5：角色成长系统——HP/攻击/防御属性，物品收集，战斗经验获取")

doc.add_heading("二、系统架构设计", level=1)

doc.add_heading("1. 整体架构图", level=2)
doc.add_paragraph(
    "系统采用分层架构设计，自上而下分为：\n"
    "- UI层（Streamlit）：负责用户交互、界面展示和状态管理\n"
    "- 编排层（LangChain LCEL）：使用管道符|串联Prompt、LLM、工具调用等组件\n"
    "- 工具层（Tools）：探索、战斗、线索、状态查询、休息等自定义工具\n"
    "- 记忆层：短期记忆（对话历史列表）+ 长期记忆（ChromaDB向量数据库）\n"
    "- 数据层：SQLite存储用户信息、游戏状态、聊天记录、线索；ChromaDB存储世界观知识\n"
    "- 模型层：OpenAI GPT-4o-mini作为核心LLM"
)

doc.add_heading("2. 核心组件选型", level=2)
doc.add_paragraph("大语言模型：OpenAI GPT-4o-mini（支持Function Calling，成本低且响应快）")
doc.add_paragraph("向量数据库：ChromaDB（轻量级，支持本地持久化，适合学生项目）")
doc.add_paragraph("关系型数据库：SQLite（无需额外部署，适合本地开发和演示）")

doc.add_heading("三、核心技术实现", level=1)

doc.add_heading("1. LCEL 与 Runnable 编排", level=2)
doc.add_paragraph(
    "项目全面采用LangChain Expression Language (LCEL)进行工作流编排，核心链路如下：\n\n"
    "chain = prompt | llm_with_tools\n\n"
    "其中prompt为ChatPromptTemplate，包含系统提示词、世界知识上下文和对话历史三个部分；"
    "llm_with_tools为绑定了工具列表的ChatOpenAI实例。当LLM决定调用工具时，系统执行tool_calls，"
    "将工具返回结果注入下一轮Prompt，再次调用LLM生成最终响应。"
    "这种管道式编排使得代码逻辑清晰，易于维护和扩展。"
)

doc.add_heading("2. 记忆机制(Memory)策略", level=2)
doc.add_paragraph(
    "短期记忆：使用Python列表模拟ConversationBufferWindowMemory，"
    "通过db_manager.get_chat_history()获取最近20条对话记录，格式化为HumanMessage/AIMessage后注入Prompt的MessagesPlaceholder。"
    "这确保了AI在每轮对话中都能理解上下文。\n\n"
    "长期记忆：使用ChromaDB向量数据库存储世界观设定、角色背景、剧情线索等持久化知识。"
    "每次用户输入时，通过similarity_search检索最相关的3条知识片段，注入到system prompt的【世界知识】部分。"
    "这样即使对话轮数很多，AI也不会遗忘核心世界观设定。\n\n"
    "线索持久化：通过SQLite数据库的world_clues表存储玩家发现的所有线索，"
    "支持跨会话的知识积累。玩家可以随时使用「查看线索」命令回顾已发现的信息。"
)

doc.add_heading("3. 工具调用(Tools)与Agent逻辑", level=2)
doc.add_paragraph(
    "项目定义了5个自定义工具，均使用@tool装饰器定义，符合LangChain Tool规范：\n\n"
    "1. push_plot（推进剧情）：根据输入的方向（探索/对话/战斗/逃离）生成对应的剧情片段，"
    "同时有概率触发线索发现，更新角色HP和章节进度。\n"
    "2. get_clues（获取线索）：从SQLite数据库中读取当前会话已收集的所有线索并展示。\n"
    "3. battle（战斗）：模拟回合制战斗，根据角色属性和敌人属性计算伤害，"
    "战斗胜利后有概率获得物品奖励。\n"
    "4. check_status（查看状态）：返回角色当前的HP、攻击力、防御力、物品栏等信息。\n"
    "5. rest（休息）：恢复一定量的HP，确保玩家不会因生命值过低而卡关。\n\n"
    "Agent逻辑：LLM接收到用户输入后，自主判断是否需要调用工具。当用户说「探索」时，"
    "LLM会自动调用push_plot工具；当用户询问「我有什么线索」时，会调用get_clues工具。"
    "这种「思考-行动-观察」循环使得游戏体验更加自然流畅。"
)

doc.add_heading("4. 数据库交互设计", level=2)
doc.add_paragraph(
    "SQLite数据库包含以下4张表：\n\n"
    "1. users表：存储用户基本信息（user_id, username, created_at）\n"
    "2. game_sessions表：存储游戏会话状态，包括当前章节、场景、HP、攻击、防御、物品栏（JSON格式）、剧情标记等\n"
    "3. chat_history表：存储对话历史，支持多轮对话记忆\n"
    "4. world_clues表：存储玩家发现的剧情线索，支持跨会话积累\n\n"
    "数据库交互通过db_manager.py模块封装，提供create_user、create_session、get_session、"
    "update_session、save_chat、get_chat_history、save_clue、get_clues等方法，"
    "使用上下文管理器确保数据库连接的正确关闭。"
)

doc.add_heading("四、交互界面(UI)设计", level=1)

doc.add_heading("1. 界面布局与功能", level=2)
doc.add_paragraph(
    "Streamlit界面采用左侧边栏+右侧主聊天区的布局：\n\n"
    "左侧边栏：\n"
    "- 世界选择区：提供「暗影纪元」和「仙途」两个按钮，点击即可开始新游戏\n"
    "- 角色状态区：实时显示HP、章节、攻防属性、物品栏\n"
    "- 线索展示区：列出已收集的所有线索\n\n"
    "右侧主聊天区：\n"
    "- 标准聊天界面，支持消息气泡展示\n"
    "- 底部输入框，支持自由文本输入\n"
    "- AI响应支持Markdown格式渲染，增强可读性"
)

doc.add_heading("2. 状态管理(Session State)", level=2)
doc.add_paragraph(
    "使用st.session_state解决Streamlit页面刷新导致状态丢失的问题：\n\n"
    "- user_id：唯一用户标识，自动生成\n"
    "- session_id：当前游戏会话ID，对应SQLite中的game_sessions表\n"
    "- messages：前端展示的对话历史列表\n"
    "- novel_title：当前选择的世界主题\n"
    "- game_chain：当前游戏的LCEL处理链实例\n\n"
    "所有持久化数据（游戏状态、对话历史、线索）均存储在SQLite数据库中，"
    "确保即使页面刷新，游戏进度也不会丢失。"
)

doc.add_heading("五、项目难点与解决方案", level=1)

doc.add_heading("1. 遇到的技术挑战", level=2)
doc.add_paragraph(
    "1. Agent工具调用循环问题：初期LLM在调用工具后容易陷入重复调用的死循环，"
    "导致响应缓慢且内容重复。\n"
    "2. 世界观记忆遗忘：随着对话轮数增加，LLM容易遗忘核心世界观设定，"
    "输出偏离预设风格。\n"
    "3. 工具参数传递：Streamlit会话状态与工具函数之间的参数传递需要正确处理session_id。"
)

doc.add_heading("2. 解决思路与优化过程", level=2)
doc.add_paragraph(
    "1. 限制工具调用轮数：通过在invoke_chain函数中设置逻辑，确保工具最多调用一次。"
    "当检测到tool_calls时，执行工具后将结果注入下一轮Prompt，由LLM生成最终响应，避免循环。\n"
    "2. RAG增强记忆：使用ChromaDB向量数据库存储世界观知识，每次对话前通过similarity_search"
    "检索最相关的3条知识片段注入Prompt，确保AI始终记住核心设定。\n"
    "3. Session ID绑定：在工具函数中添加session_id参数，通过tool_args[\"session_id\"] = session_id"
    "的方式将当前会话ID传递给工具，确保工具能正确读写数据库。\n"
    "4. 异常处理：在Streamlit调用链中添加try-except，捕获API错误并给用户友好的提示。"
)

doc.add_heading("六、总结与未来展望", level=1)

doc.add_heading("1. 项目总结", level=2)
doc.add_paragraph(
    "通过本次结业项目，深入理解了LangChain框架的核心概念和工程化实践。"
    "掌握了LCEL管道式编排的优雅写法，体验了工具调用（Function Calling）的Agent设计模式，"
    "理解了长短期记忆结合在实际应用中的重要性。项目从需求分析、架构设计到代码实现，"
    "完整经历了一个AI应用的开发流程，对RAG、Agent、Memory等LangChain核心模块有了实践层面的深入理解。"
)

doc.add_heading("2. 优化与拓展方向", level=2)
doc.add_paragraph(
    "1. 多模态支持：集成图片生成能力，为每个场景生成配图，增强沉浸感\n"
    "2. 多人协作：支持多个玩家在同一世界中互动，实现真正的多人互动小说\n"
    "3. 存档系统：支持游戏存档和读档，允许玩家在不同分支点重新选择\n"
    "4. 更丰富的世界观：扩展更多主题世界（如赛博朋克、末日生存等）\n"
    "5. 语音交互：集成TTS/STT技术，实现语音对话\n"
    "6. 成就系统：设计成就和任务系统，增加游戏目标感"
)

doc.add_heading("附录", level=1)

doc.add_heading("环境依赖", level=2)
doc.add_paragraph("requirements.txt内容：")
doc.add_paragraph(
    "langchain>=0.3.0\n"
    "langchain-core>=0.3.0\n"
    "langchain-community>=0.3.0\n"
    "langchain-openai>=0.2.0\n"
    "openai>=1.0.0\n"
    "chromadb>=0.5.0\n"
    "streamlit>=1.30.0\n"
    "python-dotenv>=1.0.0\n"
    "tiktoken>=0.5.0"
)

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "实验报告-方向四-沉浸式多角色互动小说引擎.docx")
doc.save(output_path)
print(f"报告已生成: {output_path}")
